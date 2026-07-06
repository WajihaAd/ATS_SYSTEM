"""
Resume Extractor (Gemini + PostgreSQL) - FIXED PRODUCTION VERSION
------------------------------------------------------------------
✔ PDF / DOCX / TXT extraction
✔ Robust DOCX table + paragraph parsing
✔ Gemini structured extraction (stable prompt)
✔ PostgreSQL insert with dedup (file_hash)
✔ Safe retry + error handling
✔ Debug-friendly logging
"""

import os
import json
import hashlib
import tempfile
import random
import time

import psycopg
import pdfplumber
import docx

from dotenv import load_dotenv
from google import genai
from google.genai.errors import ServerError

# =========================
# ENV
# =========================
load_dotenv()

API_KEY = os.getenv("GEMINI_API_KEY")
DB_PASSWORD = os.getenv("DB_PASSWORD")

if not API_KEY:
    raise ValueError("Missing GEMINI_API_KEY")

if not DB_PASSWORD:
    raise ValueError("Missing DB_PASSWORD")

client = genai.Client(api_key=API_KEY)

DEFAULT_MODEL = "gemini-2.5-flash"

print("🚀 Resume extractor starting...")

# =========================
# DB
# =========================
def get_db_connection():
    conn = psycopg.connect(
        dbname="resume_db",
        user="postgres",
        password=DB_PASSWORD,
        host="localhost",
        port="5432"
    )
    print("📦 DB Connected:", conn.info.dbname)
    return conn

# =========================
# HASH
# =========================
def get_resume_hash(file_bytes: bytes) -> str:
    return hashlib.md5(file_bytes).hexdigest()

# =========================
# FILE EXTRACTION (FIXED)
# =========================
def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().split(".")[-1]

    # -------- PDF --------
    if ext == "pdf":
        text = []
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            path = tmp.name

        try:
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text.append(t)
        finally:
            os.unlink(path)

        return "\n".join(text)

    # -------- DOCX (FIXED STRONGLY) --------
    elif ext == "docx":
        text = []

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file_bytes)
            path = tmp.name

        try:
            d = docx.Document(path)

            # paragraphs
            for p in d.paragraphs:
                if p.text.strip():
                    text.append(p.text.strip())

            # tables (VERY IMPORTANT for resumes)
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text.strip())

        finally:
            os.unlink(path)

        return "\n".join(text)

    # -------- TXT --------
    elif ext in ("txt", "md"):
        return file_bytes.decode("utf-8", errors="ignore")

    else:
        raise ValueError(f"Unsupported file type: {ext}")

# =========================
# GEMINI PROMPT (FIXED)
# =========================
SYSTEM_PROMPT = """
You are an expert ATS resume parser.

Extract ALL possible structured information from the resume.

Rules:
- Do NOT invent data
- Use "" for missing strings
- Use [] for missing lists
- Use 0 for missing numbers
- Be aggressive in extraction (skills, companies, education, roles)
- Output ONLY valid JSON
"""

JSON_SCHEMA = {
    "name": "",
    "email": "",
    "phone": "",
    "location": "",
    "education": [
        {
            "degree": "",
            "major": "",
            "university": "",
            "cgpa": "",
            "graduation_year": ""
        }
    ],
    "skills": [],
    "certifications": [],
    "languages": [],
    "leadership": [],
    "achievements": [],
    "consultancy_companies": [],
    "mega_projects": [],
    "total_experience_years": 0
}

# =========================
# GEMINI CALL (FIXED)
# =========================
def call_gemini(text: str, model: str = DEFAULT_MODEL) -> dict:

    text = " ".join(text.split())

    if len(text) < 50:
        raise ValueError("❌ Extracted text is too small — file parsing failed")

    text = text[:7000]

    prompt = f"""
{SYSTEM_PROMPT}

SCHEMA:
{json.dumps(JSON_SCHEMA, indent=2)}

RESUME:
{text}
"""

    last_error = None

    for attempt in range(5):
        try:
            time.sleep(random.uniform(0.5, 1.2))

            response = client.models.generate_content(
                model=model,
                contents=prompt,
                config={
                    "response_mime_type": "application/json",
                    "temperature": 0
                }
            )

            raw = response.text.strip()

            if raw.startswith("```"):
                raw = raw.strip("`")

            return json.loads(raw)

        except ServerError as e:
            last_error = e
            time.sleep(2 * (attempt + 1))

        except json.JSONDecodeError:
            print("⚠️ RAW GEMINI OUTPUT:", raw[:500])
            raise ValueError("Gemini returned invalid JSON")

    raise RuntimeError(f"Gemini failed: {last_error}")

# =========================
# NORMALIZATION
# =========================
def normalize(raw: dict, filename: str) -> dict:

    edu = (raw.get("education") or [{}])[0]

    return {
        "name": raw.get("name") or "Unknown",
        "email": raw.get("email") or "NULL",
        "phone": raw.get("phone") or "NULL",
        "location": raw.get("location") or "NULL",

        "degree": edu.get("degree") or "NULL",
        "major": edu.get("major") or "NULL",
        "cgpa": edu.get("cgpa") or "NULL",
        "university": edu.get("university") or "NULL",
        "graduation_year": edu.get("graduation_year") or "NULL",

        "experience_years": raw.get("total_experience_years", 0),

        "skills": raw.get("skills") or [],
        "certifications": raw.get("certifications") or [],
        "languages": raw.get("languages") or [],
        "leadership": raw.get("leadership") or [],
        "achievements": raw.get("achievements") or [],
        "consultancy_companies": raw.get("consultancy_companies") or [],
        "mega_projects": raw.get("mega_projects") or [],

        "filename": filename,
        "_raw": raw
    }

# =========================
# SAVE TO POSTGRES (FIXED)
# =========================
def save_to_postgres(candidate, filename, file_hash):

    conn = get_db_connection()

    try:
        with conn.cursor() as cur:

            print("💾 INSERTING:", candidate["name"])

            cur.execute("""
                INSERT INTO resumes (
                    name, email, phone, location,
                    degree, major, cgpa, university, graduation_year,
                    experience_years,
                    skills, certifications, leadership, languages,
                    consultancy_companies, mega_projects,
                    raw_data,
                    filename,
                    file_hash
                )
                VALUES (
                    %s, %s, %s, %s,
                    %s, %s, %s, %s, %s,
                    %s,
                    %s, %s, %s, %s,
                    %s, %s,
                    %s,
                    %s,
                    %s
                )
                ON CONFLICT (file_hash) DO UPDATE SET
                    name = EXCLUDED.name,
                    email = EXCLUDED.email,
                    phone = EXCLUDED.phone,
                    location = EXCLUDED.location,
                    degree = EXCLUDED.degree,
                    major = EXCLUDED.major,
                    cgpa = EXCLUDED.cgpa,
                    university = EXCLUDED.university,
                    graduation_year = EXCLUDED.graduation_year,
                    experience_years = EXCLUDED.experience_years,
                    skills = EXCLUDED.skills,
                    certifications = EXCLUDED.certifications,
                    leadership = EXCLUDED.leadership,
                    languages = EXCLUDED.languages,
                    consultancy_companies = EXCLUDED.consultancy_companies,
                    mega_projects = EXCLUDED.mega_projects,
                    raw_data = EXCLUDED.raw_data,
                    filename = EXCLUDED.filename;
            """, (
                candidate["name"],
                candidate["email"],
                candidate["phone"],
                candidate["location"],
                candidate["degree"],
                candidate["major"],
                candidate["cgpa"],
                candidate["university"],
                candidate["graduation_year"],
                candidate["experience_years"],
                json.dumps(candidate["skills"]),
                json.dumps(candidate["certifications"]),
                json.dumps(candidate["leadership"]),
                json.dumps(candidate["languages"]),
                json.dumps(candidate["consultancy_companies"]),
                json.dumps(candidate["mega_projects"]),
                json.dumps(candidate["_raw"]),
                filename,
                file_hash
            ))

        conn.commit()
        print("Rows affected:", cur.rowcount)

    except Exception as e:
        print("❌ DB ERROR:", e)

    finally:
        conn.close()

# =========================
# MAIN PIPELINE
# =========================
def parse_resume(file_bytes: bytes, filename: str, model: str = DEFAULT_MODEL):

    file_hash = get_resume_hash(file_bytes)

    text = extract_text_from_file(file_bytes, filename)

    print("\n📄 EXTRACTED TEXT PREVIEW:\n", text[:500], "\n")

    raw = call_gemini(text, model)
    result = normalize(raw, filename)

    save_to_postgres(result, filename, file_hash)

    return result

# =========================
# CLI
# =========================
if __name__ == "__main__":

    import sys

    path = sys.argv[1]

    with open(path, "rb") as f:
        data = f.read()

    result = parse_resume(data, os.path.basename(path))
    print(json.dumps(result, indent=2))